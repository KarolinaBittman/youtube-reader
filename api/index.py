import os, re, io
from flask import Flask, request, jsonify, send_file, Response
from youtube_transcript_api import YouTubeTranscriptApi, TranscriptsDisabled, NoTranscriptFound
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)

# ── helpers ──────────────────────────────────────────────────────────────────

def extract_video_id(url):
    m = re.search(r'(?:v=|/v/|youtu\.be/|/embed/|/shorts/)([a-zA-Z0-9_-]{11})', url)
    return m.group(1) if m else None

def get_transcript(video_id):
    try:
        tlist = YouTubeTranscriptApi.list_transcripts(video_id)
        for lang in ['en', 'en-US', 'en-GB']:
            try:
                return tlist.find_transcript([lang]).fetch()
            except Exception:
                pass
        try:
            return tlist.find_generated_transcript(['en']).fetch()
        except Exception:
            pass
        for t in tlist:
            return t.fetch()
    except (TranscriptsDisabled, NoTranscriptFound):
        return None
    except Exception:
        return None

def build_full_text(transcript):
    return ' '.join(seg.get('text', '') for seg in transcript).replace('\n', ' ')

def summarize_text(text, num_sentences=8):
    sentences = re.split(r'(?<=[.!?])\s+', text.strip())
    if len(sentences) <= num_sentences:
        return text
    stop_words = {
        'the','a','an','and','or','but','in','on','at','to','for','of','with',
        'is','was','are','were','be','been','being','have','has','had','do',
        'does','did','will','would','could','should','may','might','shall',
        'this','that','these','those','i','you','he','she','it','we','they',
        'my','your','his','her','its','our','their','me','him','us','them',
        'so','if','as','from','by','about','up','out','not','no','just',
    }
    word_freq = {}
    for word in re.findall(r'\b\w+\b', text.lower()):
        if word not in stop_words and len(word) > 2:
            word_freq[word] = word_freq.get(word, 0) + 1
    scored = []
    for i, sentence in enumerate(sentences):
        words = re.findall(r'\b\w+\b', sentence.lower())
        score = sum(word_freq.get(w, 0) for w in words if w not in stop_words)
        scored.append((score, i, sentence))
    top = sorted(scored, key=lambda x: x[0], reverse=True)[:num_sentences]
    return ' '.join(s[2] for s in sorted(top, key=lambda x: x[1]))

def get_video_title(video_id):
    try:
        import urllib.request, json
        url = f"https://www.youtube.com/oembed?url=https://www.youtube.com/watch?v={video_id}&format=json"
        with urllib.request.urlopen(url, timeout=5) as r:
            return json.loads(r.read()).get('title', 'YouTube Video')
    except Exception:
        return 'YouTube Video'

def create_docx(title, video_url, summary, transcript):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Source: {video_url}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()
    sh = doc.add_heading('Summary', level=1)
    for run in sh.runs:
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x50)
    doc.add_paragraph(summary).paragraph_format.space_after = Pt(12)
    doc.add_paragraph()
    th = doc.add_heading('Full Transcript', level=1)
    for run in th.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    tp = doc.add_paragraph(transcript)
    tp.paragraph_format.space_after = Pt(8)
    for run in tp.runs:
        run.font.size = Pt(10.5)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

# ── HTML (inlined so no templates folder needed on Vercel) ───────────────────

HTML = r"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8"/>
  <meta name="viewport" content="width=device-width, initial-scale=1.0"/>
  <title>YouTube Reader</title>
  <style>
    *,*::before,*::after{box-sizing:border-box;margin:0;padding:0}
    body{font-family:'Segoe UI',system-ui,-apple-system,sans-serif;background:#0f0f1a;color:#e8e8f0;min-height:100vh;display:flex;flex-direction:column;align-items:center;padding:40px 20px 80px}
    .header{text-align:center;margin-bottom:40px}
    .header .logo{font-size:2.4rem;margin-bottom:8px}
    .header h1{font-size:2rem;font-weight:700;background:linear-gradient(135deg,#ff0050,#ff6b35);-webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text}
    .header p{color:#8888aa;margin-top:8px;font-size:1rem}
    .card{background:#16162a;border:1px solid #2a2a45;border-radius:16px;padding:32px;width:100%;max-width:820px;box-shadow:0 8px 40px rgba(0,0,0,.4)}
    .input-row{display:flex;gap:12px;margin-bottom:8px}
    .url-input{flex:1;background:#0f0f1a;border:1px solid #2a2a45;border-radius:10px;color:#e8e8f0;font-size:1rem;padding:14px 18px;outline:none;transition:border-color .2s}
    .url-input:focus{border-color:#ff0050}
    .url-input::placeholder{color:#44445a}
    .btn{border:none;border-radius:10px;cursor:pointer;font-size:1rem;font-weight:600;padding:14px 26px;transition:opacity .2s,transform .1s;white-space:nowrap}
    .btn:hover{opacity:.88;transform:translateY(-1px)}
    .btn:active{transform:translateY(0)}
    .btn-primary{background:linear-gradient(135deg,#ff0050,#ff6b35);color:#fff}
    .btn:disabled{opacity:.4;cursor:not-allowed;transform:none}
    .status-msg{font-size:.9rem;min-height:20px;margin-bottom:4px}
    .status-msg.loading{color:#ffaa00}
    .status-msg.error{color:#ff4466}
    .status-msg.success{color:#44cc88}
    .spinner{display:inline-block;width:14px;height:14px;border:2px solid #ffaa00;border-top-color:transparent;border-radius:50%;animation:spin .8s linear infinite;margin-right:6px;vertical-align:middle}
    @keyframes spin{to{transform:rotate(360deg)}}
    .result-area{display:none;margin-top:28px}
    .result-area.visible{display:block}
    .result-title{font-size:1.3rem;font-weight:700;color:#fff;margin-bottom:20px;padding-bottom:12px;border-bottom:1px solid #2a2a45}
    .section-label{display:flex;align-items:center;gap:8px;font-size:.78rem;font-weight:700;letter-spacing:.1em;text-transform:uppercase;margin-bottom:10px;margin-top:22px}
    .section-label.summary-label{color:#ff0050}
    .section-label.transcript-label{color:#8888cc}
    .tag{background:rgba(255,0,80,.15);border:1px solid rgba(255,0,80,.3);border-radius:4px;padding:2px 7px;font-size:.7rem}
    .tag.blue{background:rgba(100,100,220,.15);border-color:rgba(100,100,220,.3)}
    .text-box{background:#0f0f1a;border:1px solid #2a2a45;border-radius:10px;padding:18px 20px;font-size:.95rem;line-height:1.75;color:#d0d0e8;white-space:pre-wrap;word-break:break-word}
    .text-box.transcript-box{max-height:340px;overflow-y:auto}
    .text-box.transcript-box::-webkit-scrollbar{width:6px}
    .text-box.transcript-box::-webkit-scrollbar-track{background:transparent}
    .text-box.transcript-box::-webkit-scrollbar-thumb{background:#2a2a45;border-radius:3px}
    .download-bar{display:flex;align-items:center;justify-content:space-between;background:#1a1a2e;border:1px solid #2a2a45;border-radius:10px;padding:14px 20px;margin-top:24px;gap:16px;flex-wrap:wrap}
    .download-info{font-size:.9rem;color:#8888aa}
    .download-info strong{color:#e8e8f0}
    .footer{margin-top:32px;color:#44445a;font-size:.8rem;text-align:center}
  </style>
</head>
<body>
  <div class="header">
    <div class="logo">🎬</div>
    <h1>YouTube Reader</h1>
    <p>Paste any YouTube link — read the transcript &amp; download it as a Word doc</p>
  </div>
  <div class="card">
    <div class="input-row">
      <input id="urlInput" class="url-input" type="url" placeholder="https://www.youtube.com/watch?v=..." autocomplete="off"/>
      <button id="processBtn" class="btn btn-primary" onclick="processVideo()">▶ &nbsp;Read</button>
    </div>
    <p id="statusMsg" class="status-msg"></p>
    <div id="resultArea" class="result-area">
      <div id="videoTitle" class="result-title"></div>
      <div class="section-label summary-label"><span>⚡ Summary</span><span class="tag">Key points</span></div>
      <div id="summaryBox" class="text-box"></div>
      <div class="section-label transcript-label"><span>📄 Full Transcript</span><span class="tag blue">Scrollable</span></div>
      <div id="transcriptBox" class="text-box transcript-box"></div>
      <div class="download-bar">
        <div class="download-info">Ready to save? Download everything as a Word document.</div>
        <button id="downloadBtn" class="btn btn-primary" onclick="downloadDoc()">⬇ &nbsp;Download Word Doc</button>
      </div>
    </div>
  </div>
  <div class="footer">Works with any public YouTube video that has subtitles or auto-generated captions.</div>
  <script>
    let currentData=null;
    function setStatus(msg,type){const el=document.getElementById('statusMsg');el.className='status-msg '+type;el.innerHTML=msg}
    async function processVideo(){
      const url=document.getElementById('urlInput').value.trim();
      if(!url){setStatus('Please paste a YouTube URL first.','error');return}
      const btn=document.getElementById('processBtn');
      btn.disabled=true;
      document.getElementById('resultArea').classList.remove('visible');
      setStatus('<span class="spinner"></span>Fetching transcript…','loading');
      try{
        const res=await fetch('/process',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({url})});
        const data=await res.json();
        if(!res.ok){setStatus('⚠ '+(data.error||'Something went wrong.'),'error');return}
        currentData={...data,url};
        document.getElementById('videoTitle').textContent='🎥 '+data.title;
        document.getElementById('summaryBox').textContent=data.summary;
        document.getElementById('transcriptBox').textContent=data.transcript;
        document.getElementById('resultArea').classList.add('visible');
        setStatus('✓ Done! Scroll down to read or download.','success');
      }catch(e){setStatus('⚠ Something went wrong. Please try again.','error')}
      finally{btn.disabled=false}
    }
    async function downloadDoc(){
      if(!currentData)return;
      const btn=document.getElementById('downloadBtn');
      btn.disabled=true;btn.textContent='⏳  Generating…';
      try{
        const res=await fetch('/download',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify(currentData)});
        if(!res.ok){alert('Download failed. Please try again.');return}
        const blob=await res.blob();
        const a=document.createElement('a');
        a.href=URL.createObjectURL(blob);
        const disposition=res.headers.get('Content-Disposition')||'';
        const match=disposition.match(/filename="?(.+?)"?$/);
        a.download=match?match[1]:'transcript.docx';
        document.body.appendChild(a);a.click();a.remove();
      }finally{btn.disabled=false;btn.innerHTML='⬇ &nbsp;Download Word Doc'}
    }
    document.addEventListener('DOMContentLoaded',()=>{
      document.getElementById('urlInput').addEventListener('keydown',e=>{if(e.key==='Enter')processVideo()});
    });
  </script>
</body>
</html>"""

# ── routes ───────────────────────────────────────────────────────────────────

@app.route('/')
def index():
    return Response(HTML, mimetype='text/html')

@app.route('/process', methods=['POST'])
def process():
    data = request.get_json()
    url = (data or {}).get('url', '').strip()
    if not url:
        return jsonify({'error': 'Please paste a YouTube URL.'}), 400
    video_id = extract_video_id(url)
    if not video_id:
        return jsonify({'error': 'Could not find a valid YouTube video ID in that URL.'}), 400
    transcript_data = get_transcript(video_id)
    if not transcript_data:
        return jsonify({'error': 'No transcript available for this video. It may be disabled or private.'}), 404
    full_text = build_full_text(transcript_data)
    summary = summarize_text(full_text)
    title = get_video_title(video_id)
    return jsonify({'title': title, 'video_id': video_id, 'summary': summary, 'transcript': full_text})

@app.route('/download', methods=['POST'])
def download():
    data = request.get_json()
    title = data.get('title', 'YouTube Video')
    url = data.get('url', '')
    summary = data.get('summary', '')
    transcript = data.get('transcript', '')
    docx_bytes = create_docx(title, url, summary, transcript)
    safe_title = re.sub(r'[^\w\s-]', '', title)[:50].strip().replace(' ', '_')
    return send_file(
        io.BytesIO(docx_bytes),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f"{safe_title}.docx",
    )
